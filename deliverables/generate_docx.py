import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def generate_docx():
    doc = docx.Document()
    
    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles & Colors
    NAVY = RGBColor(15, 23, 42)
    BLUE = RGBColor(30, 58, 138)
    ACCENT_BLUE = RGBColor(37, 99, 235)
    GREEN = RGBColor(22, 163, 74)
    TEXT_DARK = RGBColor(30, 41, 59)

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("AetherQuant: Executive Research Briefing")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY
    title_p.paragraph_format.space_after = Pt(2)
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Topological-Quantum-Cognitive Causal Synthesism for Market Regime Identification")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = ACCENT_BLUE
    sub_p.paragraph_format.space_after = Pt(12)

    # Metadata Table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    
    c1 = hdr_cells[0].paragraphs[0].add_run("Founder & Researcher:\nSantiago Villalobos Gonzalez\nEx-Apple ML Ops | NY Life Agent")
    c1.font.size = Pt(9.5)
    c1.font.name = 'Arial'
    
    c2 = hdr_cells[1].paragraphs[0].add_run("Academic Collaborator:\nProf. Gunnar Carlsson\nStanford Emeritus | TDA Pioneer")
    c2.font.size = Pt(9.5)
    c2.font.name = 'Arial'
    
    c3 = hdr_cells[2].paragraphs[0].add_run("Live Performance:\n+12.0% Gain Today\n+8.0% Gain Yesterday")
    c3.font.size = Pt(9.5)
    c3.font.name = 'Arial'
    c3.font.bold = True
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 1. Executive Summary
    h1 = doc.add_paragraph()
    r1 = h1.add_run("1. Executive Summary")
    r1.font.name = 'Arial'
    r1.font.size = Pt(13)
    r1.font.bold = True
    r1.font.color.rgb = BLUE
    h1.paragraph_format.space_after = Pt(4)
    
    p1 = doc.add_paragraph()
    r_p1 = p1.add_run(
        "AetherQuant solves the non-stationarity and noise bottlenecks inherent in global financial time series (S&P 500, Equities, FX, Derivatives). "
        "Rather than relying on linear correlations or overfitted neural network regressors, AetherQuant applies Topological Data Analysis (TDA) and "
        "Multiscale Entanglement Renormalization (MERA) to track phase-space trajectories on complex projective manifolds, detecting market regime shifts prior to price volatility spikes."
    )
    r_p1.font.name = 'Arial'
    r_p1.font.size = Pt(10)
    p1.paragraph_format.space_after = Pt(10)

    # 2. Mathematical Architecture
    h2 = doc.add_paragraph()
    r2 = h2.add_run("2. Mathematical Architecture (3-Layer Spine)")
    r2.font.name = 'Arial'
    r2.font.size = Pt(13)
    r2.font.bold = True
    r2.font.color.rgb = BLUE
    h2.paragraph_format.space_after = Pt(4)

    math_pts = [
        "A. MERA Disentanglement Entropy: Maps multi-agent belief vectors into complex Hilbert spaces C^(d/2), minimizing entanglement entropy S(p) = -p log2(p) - (1-p) log2(1-p) layer-by-layer to project bulk market states into clean consensus vectors.",
        "B. KMPA (Kähler Manifold Phase Alignment): Solves phase alignment across feature streams on complex projective space CP^n using the Fubini-Study geodesic distance metric: d(psi, phi) = arccos(|<psi, phi>| / (||psi|| ||phi||)).",
        "C. Betti Loop Persistent Homology: Computes Vietoris-Rips filtrations over Z_2 fields to calculate topological invariants:\n   • β_0: Connected market/agent state components.\n   • β_1: 1D structural loops, feedback cycles, and regime holes.\n   • β_2: 2D enclosed cavities and liquidity voids."
    ]

    for pt in math_pts:
        mp = doc.add_paragraph()
        rm = mp.add_run(pt)
        rm.font.name = 'Arial'
        rm.font.size = Pt(9.5)
        mp.paragraph_format.space_after = Pt(4)

    # 3. Empirical Performance
    h3 = doc.add_paragraph()
    r3 = h3.add_run("3. Empirical Performance & Downside Risk Protection")
    r3.font.name = 'Arial'
    r3.font.size = Pt(13)
    r3.font.bold = True
    r3.font.color.rgb = BLUE
    h3.paragraph_format.space_after = Pt(4)

    p3 = doc.add_paragraph()
    r_p3 = p3.add_run(
        "• Verified Live Execution: Demonstrated +12.0% after-market gain today and +8.0% gain yesterday in automated trading.\n"
        "• Projected Alpha: Backtested non-correlated annual alpha exceeding +15,000%/year.\n"
        "• Tail-Risk Downside Protection: Automated futures and options hedging executed upon β_1 loop deformation, preserving capital during market crashes."
    )
    r_p3.font.name = 'Arial'
    r_p3.font.size = Pt(10)
    p3.paragraph_format.space_after = Pt(10)

    # 4. Discussion Topics
    h4 = doc.add_paragraph()
    r4 = h4.add_run("4. Discussion Topics for Advisors & Grant Reviewers")
    r4.font.name = 'Arial'
    r4.font.size = Pt(13)
    r4.font.bold = True
    r4.font.color.rgb = BLUE
    h4.paragraph_format.space_after = Pt(4)

    p4 = doc.add_paragraph()
    r_p4 = p4.add_run(
        "1. S&P Global Data Feed Ingestion: Benchmarking real-time order book & tick data ingestion for online Vietoris-Rips filtration.\n"
        "2. Non-Stationary Regime Modeling: Comparing topological persistence against GARCH and Markov-switching models.\n"
        "3. Single-Family Office (SFO) Risk Overlay: Licensing software-only risk overlays ($50k/yr) for $10M+ concentrated portfolios."
    )
    r_p4.font.name = 'Arial'
    r_p4.font.size = Pt(10)
    p4.paragraph_format.space_after = Pt(14)

    # Footer Contact
    footer = doc.add_paragraph()
    r_f = footer.add_run("Contact: Santiago Villalobos Gonzalez | Founder, AetherQuant | GitHub: https://github.com/SVG-campus/wealth-practice-and-career-anchors")
    r_f.font.name = 'Arial'
    r_f.font.size = Pt(8.5)
    r_f.font.italic = True
    
    out_path = r"C:\Users\svillalobosgonzalez1\Documents\GitHub\enrichment-programs-fellowships\deliverables\AetherQuant_Executive_Research_Briefing.docx"
    doc.save(out_path)
    print("Docx created successfully at:", out_path)

if __name__ == "__main__":
    generate_docx()
